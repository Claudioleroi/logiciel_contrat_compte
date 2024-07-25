import os
import sys
import shutil
import datetime
from pathlib import Path
from PyQt5.QtWidgets import (
    QWidget, QFormLayout, QLineEdit, QPushButton, QFileDialog, QMessageBox, QTextBrowser, QVBoxLayout
)
from PyQt5.QtCore import pyqtSignal
from docxtpl import DocxTemplate
from docx import Document
from fpdf import FPDF  # For PDF export
import data.database as database

class FactureTab(QWidget):
    facture_generated = pyqtSignal()  # Signal for facture generation

    def __init__(self):
        super().__init__()
        self.init_ui()
        self.create_facture_folder()

    def init_ui(self):
        layout = QVBoxLayout()

        # Create a form layout for inputs
        form_layout = QFormLayout()

        # Define widgets
        self.client_name = QLineEdit()
        self.vendor_name = QLineEdit()
        self.amount = QLineEdit()
        self.description1 = QLineEdit()
        self.description2 = QLineEdit()

        # Add widgets to form layout
        form_layout.addRow("Client name:", self.client_name)
        form_layout.addRow("Vendor name:", self.vendor_name)
        form_layout.addRow("Amount:", self.amount)
        form_layout.addRow("Description1:", self.description1)
        form_layout.addRow("Description2:", self.description2)

        self.select_template_button = QPushButton("Select Invoice Template")
        self.select_template_button.clicked.connect(self.select_template)
        form_layout.addWidget(self.select_template_button)

        self.generate_button = QPushButton("Generate Invoice")
        self.generate_button.clicked.connect(self.generate_facture)
        form_layout.addWidget(self.generate_button)

        self.preview_button = QPushButton("Preview Invoice")
        self.preview_button.clicked.connect(self.preview_facture)
        form_layout.addWidget(self.preview_button)

        self.open_button = QPushButton("Open in Word")
        self.open_button.clicked.connect(self.open_in_word)
        form_layout.addWidget(self.open_button)

        self.print_button = QPushButton("Print Invoice")
        self.print_button.clicked.connect(self.print_facture)
        form_layout.addWidget(self.print_button)

        self.save_pdf_button = QPushButton("Save as PDF")
        self.save_pdf_button.clicked.connect(self.save_as_pdf)
        form_layout.addWidget(self.save_pdf_button)

        # Add the form layout to the main layout
        layout.addLayout(form_layout)

        # Create a text browser for template preview
        self.preview_browser = QTextBrowser()
        layout.addWidget(self.preview_browser)

        self.setLayout(layout)
        self.setWindowTitle('Invoice Generator')
        self.resize(800, 600)  # Set the window size (width, height)

        # Initialize template path and output path
        self.template_path = None
        self.temp_template_path = None
        self.output_path = None

    def create_facture_folder(self):
        """
        Create a folder named 'Facture_new' in the user's Downloads directory to store generated invoices.
        """
        if sys.platform == "win32":
            self.downloads_folder = Path(os.getenv('USERPROFILE')) / 'Downloads'
        elif sys.platform == "darwin":
            self.downloads_folder = Path.home() / 'Downloads'
        else:  # Linux and other Unix-like OS
            self.downloads_folder = Path.home() / 'Téléchargements'

        self.facture_folder = self.downloads_folder / 'Facture_new'
        self.facture_folder.mkdir(exist_ok=True)

    def select_template(self):
        """
        Open a file dialog to select an invoice template in Word format (.docx).
        Display the selected template's content in the preview browser.
        """
        file_name, _ = QFileDialog.getOpenFileName(self, "Select Invoice Template", "", "Word Documents (*.docx)")
        if file_name:
            self.template_path = Path(file_name)
            # Display a preview of the selected template
            self.preview_template()
            QMessageBox.information(self, 'Template Selected', f"Template selected: {self.template_path}")

    def preview_template(self):
        """
        Display the content of the selected template in the preview browser.
        """
        if not self.template_path:
            return

        try:
            doc = Document(self.template_path)
            doc_content = "\n".join(p.text for p in doc.paragraphs)
            self.preview_browser.setPlainText(doc_content)
        except Exception as e:
            QMessageBox.critical(self, 'Error', f"An error occurred while previewing the template: {e}")

    def generate_facture(self):
        """
        Generate an invoice by filling the selected template with user-provided data.
        Save the generated invoice to the 'Facture_new' folder and insert the contract into the database.
        """
        if not self.template_path:
            QMessageBox.critical(self, 'Error', 'No template selected.')
            return

        # Retrieve user input values
        values = {
            "CLIENT": self.client_name.text(),
            "VENDOR": self.vendor_name.text(),
            "AMOUNT": self.amount.text(),
            "LINE1": self.description1.text(),
            "LINE2": self.description2.text(),
        }

        # Add calculated fields to our dict
        try:
            amount = float(values["AMOUNT"])
            values["NONREFUNDABLE"] = round(amount * 0.2, 2)
        except ValueError:
            QMessageBox.critical(self, 'Error', 'Invalid amount value.')
            return

        # Add current date to values
        today = datetime.datetime.today()
        values["TODAY"] = today.strftime("%Y-%m-%d")

        # Copy the selected template to a temporary path
        self.temp_template_path = Path(__file__).parent / "temp_template.docx"
        shutil.copy(self.template_path, self.temp_template_path)

        # Load template and save new document
        try:
            doc = DocxTemplate(self.temp_template_path)
            doc.render(values)

            # Save the filled document to the "Facture_new" folder
            self.output_path = self.facture_folder / f"{values['VENDOR']}-invoice.docx"
            doc.save(self.output_path)

            # Optionally, delete the temporary template
            self.temp_template_path.unlink()

            # Insert contract into database
            database.insert_contract(values["CLIENT"], values["VENDOR"], values["AMOUNT"], values["LINE1"], values["LINE2"], values["TODAY"])

            # Emit signal that a contract has been generated
            self.facture_generated.emit()

            QMessageBox.information(self, 'Success', f"File has been saved here: {self.output_path}")

            # Preview the generated invoice
            self.preview_facture()
        except Exception as e:
            QMessageBox.critical(self, 'Error', f"An error occurred: {e}")

    def preview_facture(self):
        """
        Display the content of the generated invoice in the preview browser.
        """
        if not self.output_path:
            QMessageBox.critical(self, 'Error', 'No invoice created.')
            return

        try:
            doc = Document(self.output_path)
            doc_content = "\n".join(p.text for p in doc.paragraphs)
            self.preview_browser.setPlainText(doc_content)
        except Exception as e:
            QMessageBox.critical(self, 'Error', f"An error occurred while previewing the invoice: {e}")

    def open_in_word(self):
        """
        Open the generated invoice in Microsoft Word or the default document editor.
        """
        if not self.output_path:
            QMessageBox.critical(self, 'Error', 'No invoice created.')
            return

        try:
            if sys.platform == "win32":
                os.startfile(self.output_path)
            elif sys.platform == "darwin":
                os.system(f'open "{self.output_path}"')
            else:  # Linux and other Unix-like OS
                os.system(f'xdg-open "{self.output_path}"')
        except Exception as e:
            QMessageBox.critical(self, 'Error', f"An error occurred while opening the file: {e}")

    def print_facture(self):
        """
        Send the generated invoice to the printer using the default printing command.
        """
        if not self.output_path:
            QMessageBox.critical(self, 'Error', 'No invoice created.')
            return

        try:
            if sys.platform == "win32":
                # Windows command for printing
                os.system(f'start /MIN cmd /c "print /d:PRINTERNAME "{self.output_path}""')
            elif sys.platform == "darwin":
                # macOS command for printing
                os.system(f'lpr "{self.output_path}"')
            else:  # Linux and other Unix-like OS
                # Linux command for printing
                os.system(f'lpr "{self.output_path}"')
        except Exception as e:
            QMessageBox.critical(self, 'Error', f"An error occurred while printing the file: {e}")

    def save_as_pdf(self):
        """
        Convert the generated invoice from Word format to PDF and save it in the 'Facture_new' folder.
        """
        if not self.output_path:
            QMessageBox.critical(self, 'Error', 'No invoice created.')
            return

        try:
            # Ensure PyMuPDF (fitz) can be used to convert DOCX to PDF by creating a PDF directly
            pdf_path = self.facture_folder / f"{Path(self.output_path).stem}.pdf"

            # Using Word to PDF conversion for accuracy
            if sys.platform == "win32":
                # For Windows, you can use Microsoft Word COM object
                import comtypes.client

                # Initialize Word application
                word = comtypes.client.CreateObject('Word.Application')
                doc = word.Documents.Open(str(self.output_path))

                # Export as PDF
                doc.SaveAs(str(pdf_path), FileFormat=17)  # 17 corresponds to wdFormatPDF

                # Close the document and Word application
                doc.Close()
                word.Quit()

            elif sys.platform == "darwin":
                # On macOS, we can use a command line utility to convert docx to PDF
                os.system(f"libreoffice --headless --convert-to pdf --outdir {self.facture_folder} {self.output_path}")

            else:
                # For Linux, ensure libreoffice is installed for conversion
                os.system(f"libreoffice --headless --convert-to pdf --outdir {self.facture_folder} {self.output_path}")

            QMessageBox.information(self, 'Success', f"File saved as PDF here: {pdf_path}")

        except Exception as e:
            QMessageBox.critical(self, 'Error', f"An error occurred while saving as PDF: {e}")
