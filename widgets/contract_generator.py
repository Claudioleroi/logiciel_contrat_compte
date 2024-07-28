import os
import sys
import datetime
import shutil
from pathlib import Path
from PyQt5.QtWidgets import (
    QWidget, QFormLayout, QLineEdit, QPushButton, QFileDialog, QMessageBox, QTextBrowser, QVBoxLayout
)
from PyQt5.QtCore import pyqtSignal
from docxtpl import DocxTemplate
from docx import Document
import data.database as database

class ContractGenerator(QWidget):
    contract_generated = pyqtSignal()  # Signal for contract generation

    def __init__(self):
        super().__init__()
        self.init_ui()
        self.create_contract_folder()

    def init_ui(self):
        # Define layout
        layout = QVBoxLayout()

        # Create a form layout for inputs
        form_layout = QFormLayout()
        
        # Define widgets
        self.client_name = QLineEdit()
        self.vendor_name = QLineEdit()
        self.amount = QLineEdit()
        self.description1 = QLineEdit()
        self.description2 = QLineEdit()

        # Add widgets to form layout
        form_layout.addRow("Client name:", self.client_name)
        form_layout.addRow("Vendor name:", self.vendor_name)
        form_layout.addRow("Amount:", self.amount)
        form_layout.addRow("Description1:", self.description1)
        form_layout.addRow("Description2:", self.description2)

        self.select_template_button = QPushButton("Select Contract Template")
        self.select_template_button.clicked.connect(self.select_template)
        form_layout.addWidget(self.select_template_button)

        self.generate_button = QPushButton("Generate Contract")
        self.generate_button.clicked.connect(self.generate_contract)
        form_layout.addWidget(self.generate_button)

        self.preview_button = QPushButton("Preview Contract")
        self.preview_button.clicked.connect(self.preview_contract)
        form_layout.addWidget(self.preview_button)

        self.open_button = QPushButton("Open in Word")
        self.open_button.clicked.connect(self.open_in_word)
        form_layout.addWidget(self.open_button)

        self.print_button = QPushButton("Print Contract")
        self.print_button.clicked.connect(self.print_contract)
        form_layout.addWidget(self.print_button)
        
        # Add the form layout to the main layout
        layout.addLayout(form_layout)

        # Create a text browser for template preview
        self.preview_browser = QTextBrowser()
        layout.addWidget(self.preview_browser)

        self.setLayout(layout)
        self.setWindowTitle('Contract Generator')
        self.resize(800, 600)  # Set the window size (width, height)

        # Initialize template path and output path
        self.template_path = None
        self.temp_template_path = None
        self.output_path = None

    def create_contract_folder(self):
        # Determine the Downloads directory path
        if sys.platform == "win32":
            self.downloads_folder = Path(os.getenv('USERPROFILE')) / 'Downloads'
        elif sys.platform == "darwin":
            self.downloads_folder = Path.home() / 'Downloads'
        else:  # Linux and other Unix-like OS
            self.downloads_folder = Path.home() / 'Téléchargements'
        
        # Create "nouveau_contrat" folder in Downloads if it doesn't exist
        self.contract_folder = self.downloads_folder / 'nouveau_contrat'
        self.contract_folder.mkdir(exist_ok=True)

    def select_template(self):
        # Open file dialog to select the template
        file_name, _ = QFileDialog.getOpenFileName(self, "Select Contract Template", "", "Word Documents (*.docx)")
        if file_name:
            self.template_path = Path(file_name)
            # Display a preview of the selected template
            self.preview_template()
            QMessageBox.information(self, 'Template Selected', f"Template selected: {self.template_path}")

    def preview_template(self):
        if not self.template_path:
            return
        
        try:
            # Read the content of the template (text only) using python-docx
            doc = Document(self.template_path)
            doc_content = "\n".join(p.text for p in doc.paragraphs)
            self.preview_browser.setPlainText(doc_content)
        except Exception as e:
            QMessageBox.critical(self, 'Error', f"An error occurred while previewing the template: {e}")

    def generate_contract(self):
        if not self.template_path:
            QMessageBox.critical(self, 'Error', 'No template selected.')
            return

        values = {
            "CLIENT": self.client_name.text(),
            "VENDOR": self.vendor_name.text(),
            "AMOUNT": self.amount.text(),
            "LINE1": self.description1.text(),
            "LINE2": self.description2.text(),
        }

        # Add calculated fields to our dict
        try:
            amount = float(values["AMOUNT"])
            values["NONREFUNDABLE"] = round(amount * 0.2, 2)
        except ValueError:
            QMessageBox.critical(self, 'Error', 'Invalid amount value.')
            return

        today = datetime.datetime.today()
        values["TODAY"] = today.strftime("%Y-%m-%d")
        values["TODAY_IN_ONE_WEEK"] = (today + datetime.timedelta(days=7)).strftime("%Y-%m-%d")

        # Copy the selected template to a temporary path
        self.temp_template_path = Path(__file__).parent / "temp_template.docx"
        shutil.copy(self.template_path, self.temp_template_path)

        # Load template and save new document
        try:
            doc = DocxTemplate(self.temp_template_path)
            doc.render(values)
            
            # Save the filled document to the "nouveau_contrat" folder
            self.output_path = self.contract_folder / f"{values['VENDOR']}-contract.docx"
            doc.save(self.output_path)

            # Optionally, delete the temporary template
            self.temp_template_path.unlink()

            # Insert contract into database
            database.insert_contract(values["CLIENT"], values["VENDOR"], values["AMOUNT"], values["LINE1"], values["LINE2"], values["TODAY"])

            # Emit signal that a contract has been generated
            self.contract_generated.emit()

            QMessageBox.information(self, 'Success', f"File has been saved here: {self.output_path}")

            # Preview the generated contract
            self.preview_contract()

            # Clear form fields after generating the contract
            self.clear_form_fields()
        except Exception as e:
            QMessageBox.critical(self, 'Error', f"An error occurred: {e}")

    def clear_form_fields(self):
        self.client_name.clear()
        self.vendor_name.clear()
        self.amount.clear()
        self.description1.clear()
        self.description2.clear()

    def preview_contract(self):
        if not self.output_path:
            QMessageBox.critical(self, 'Error', 'No contract created.')
            return
        
        try:
            # Read the content of the modified document using python-docx
            doc = Document(self.output_path)
            doc_content = "\n".join(p.text for p in doc.paragraphs)
            self.preview_browser.setPlainText(doc_content)
        except Exception as e:
            QMessageBox.critical(self, 'Error', f"An error occurred while previewing the contract: {e}")

    def open_in_word(self):
        if not self.output_path:
            QMessageBox.critical(self, 'Error', 'No contract created.')
            return
        
        try:
            # Open the document with the default Word application
            if sys.platform == "win32":
                os.startfile(self.output_path)
            elif sys.platform == "darwin":
                os.system(f'open "{self.output_path}"')
            else:  # Linux and other Unix-like OS
                os.system(f'xdg-open "{self.output_path}"')
        except Exception as e:
            QMessageBox.critical(self, 'Error', f"An error occurred while opening the file: {e}")

    def print_contract(self):
        if not self.output_path:
            QMessageBox.critical(self, 'Error', 'No contract created.')
            return
        
        try:
            # Open the document with the default application and send it to print
            if sys.platform == "win32":
                os.system(f'print /D:printer_name "{self.output_path}"')
            elif sys.platform == "darwin":
                os.system(f'lpr "{self.output_path}"')
            else:  # Linux and other Unix-like OS
                os.system(f'lpr "{self.output_path}"')
        except Exception as e:
            QMessageBox.critical(self, 'Error', f"An error occurred while printing the file: {e}")
